#!/usr/bin/env python3
"""Build the public English academic CV from the website's structured data."""

from __future__ import annotations

import json
import subprocess
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "cv" / "Mengfei-Li-CV-English.docx"
SITE_URL = "https://murphylicn.github.io/MengfeiLi.github.io"
NAVY = "17365D"
BLUE = "2E74B5"
MUTED = "5B6573"
LIGHT = "D9E2F3"
TEXT = "1F2328"
BODY_FONT = "Arial"
SERIF_FONT = "Times New Roman"
CONTENT_WIDTH = Inches(7.2)


def read_front_matter(path: Path) -> dict:
    text = path.read_text(encoding="utf-8")
    _, front_matter, _ = text.split("---", 2)
    return parse_yaml(front_matter)


def parse_yaml(text: str):
    """Parse repository YAML with the Ruby dependency already required by Jekyll."""
    script = (
        "require 'yaml'; require 'json'; require 'date'; "
        "value = YAML.safe_load(STDIN.read, permitted_classes: [Date], aliases: true); "
        "print JSON.generate(value)"
    )
    result = subprocess.run(
        ["ruby", "-e", script],
        input=text,
        text=True,
        capture_output=True,
        check=True,
    )
    return json.loads(result.stdout)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=40, start=60, bottom=40, end=60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_bottom_border(paragraph, color=NAVY, size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)


def set_run_font(run, name=BODY_FONT, size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, *, bold=False, italic=False, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), BODY_FONT)
    r_fonts.set(qn("w:hAnsi"), BODY_FONT)
    r_pr.append(r_fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    if italic:
        r_pr.append(OxmlElement("w:i"))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, value, end):
        run._r.append(node)


def add_section_heading(doc: Document, text: str, *, page_break_before=False) -> None:
    paragraph = doc.add_paragraph(style="CV Section")
    paragraph.paragraph_format.page_break_before = page_break_before
    paragraph.add_run(text.upper())
    set_bottom_border(paragraph)


def add_compact_paragraph(doc: Document, *, left=0, hanging=0, after=2, keep=True):
    paragraph = doc.add_paragraph(style="CV Body")
    paragraph.paragraph_format.left_indent = Inches(left)
    if hanging:
        paragraph.paragraph_format.first_line_indent = Inches(-hanging)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_together = keep
    return paragraph


def add_date_entry(doc: Document, title: str, date: str, detail_lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="CV Body")
    paragraph.paragraph_format.tab_stops.add_tab_stop(CONTENT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(title)
    set_run_font(run, bold=True)
    date_run = paragraph.add_run(f"\t{date}")
    set_run_font(date_run, bold=True, color=MUTED)
    for index, line in enumerate(detail_lines):
        detail = doc.add_paragraph(style="CV Detail")
        detail.add_run(line)
        detail.paragraph_format.keep_with_next = index < len(detail_lines) - 1


def add_authors(paragraph, authors: list[str]) -> None:
    for index, author in enumerate(authors):
        run = paragraph.add_run(author)
        set_run_font(run, bold=author == "Mengfei Li")
        if index < len(authors) - 1:
            paragraph.add_run(", ")
    paragraph.add_run(". ")


def add_publication(doc: Document, publication: dict, number: int | None = None) -> None:
    paragraph = add_compact_paragraph(doc, left=0.23, hanging=0.23, after=3)
    if number is not None:
        paragraph.add_run(f"{number}. ")
    add_authors(paragraph, publication["authors"])
    publication_title = publication.get("title_en", publication["title"])
    title_url = publication.get("paper_url") or publication.get("ssrn_url")
    if title_url:
        add_hyperlink(paragraph, publication_title, title_url, italic=True, color=TEXT)
    else:
        title_run = paragraph.add_run(publication_title)
        set_run_font(title_run, name=SERIF_FONT, italic=True)
    paragraph.add_run(". ")
    status = publication["status"]
    publication_venue = publication.get("venue_en", publication.get("venue", ""))
    if status == "published":
        venue = paragraph.add_run(publication_venue)
        set_run_font(venue, italic=True)
        paragraph.add_run(f", {publication.get('pages', '')}. ")
    elif status == "forthcoming":
        paragraph.add_run("Accepted and forthcoming in ")
        venue = paragraph.add_run(publication_venue)
        set_run_font(venue, italic=True)
        paragraph.add_run(". ")
    elif status == "under-review":
        paragraph.add_run("Under second-round review at ")
        venue = paragraph.add_run(publication_venue)
        set_run_font(venue, italic=True)
        paragraph.add_run(". ")
    else:
        paragraph.add_run("Working paper. ")
    if publication.get("author_note"):
        note = paragraph.add_run(publication["author_note"] + " ")
        set_run_font(note, italic=True, color=MUTED)
    links = []
    if publication.get("ssrn_url"):
        links.append(("SSRN", publication["ssrn_url"]))
    if publication.get("doi"):
        links.append(("DOI", f"https://doi.org/{publication['doi']}"))
    if publication.get("software_url"):
        links.append(("Software", publication["software_url"]))
    for index, (label, url) in enumerate(links):
        if index:
            paragraph.add_run(" · ")
        add_hyperlink(paragraph, label, url)


def add_project(doc: Document, project: dict) -> None:
    paragraph = add_compact_paragraph(doc, left=0.23, hanging=0.23, after=3)
    role = paragraph.add_run(f"{project['role_en']}. ")
    set_run_font(role, bold=True)
    paragraph.add_run(f"{project['funder_en']}. ")
    title = paragraph.add_run(project["title_en"])
    set_run_font(title, name=SERIF_FONT, italic=True)
    date = paragraph.add_run(f". {project['display_date_en']}.")
    set_run_font(date, color=MUTED)


def add_award(doc: Document, award: dict) -> None:
    paragraph = add_compact_paragraph(doc, left=0.23, hanging=0.23, after=2)
    if award.get("url"):
        add_hyperlink(paragraph, award["title_en"], award["url"], bold=True, color=TEXT)
    else:
        run = paragraph.add_run(award["title_en"])
        set_run_font(run, bold=True)
    paragraph.add_run(f", {award['organization_en']}, {award['year']}")
    if award.get("note_en"):
        paragraph.add_run(f" ({award['note_en']})")
    paragraph.add_run(".")


def add_talks(doc: Document, talks: list[dict], publications: list[dict]) -> None:
    publication_by_project = {
        publication.get("project_id"): publication
        for publication in publications
        if publication.get("project_id")
    }
    groups: dict[str, list[dict]] = {}
    for talk in sorted(talks, key=lambda item: item["sort_date"], reverse=True):
        if talk["presentation_type"] == "Conference presentation":
            groups.setdefault(talk["paper_id"], []).append(talk)
    for project_id, items in groups.items():
        title = publication_by_project.get(project_id, {}).get("title", project_id)
        heading = doc.add_paragraph(style="CV Subsection")
        heading.add_run(title)
        for talk in items:
            paragraph = add_compact_paragraph(doc, left=0.23, hanging=0.23, after=1)
            venue = paragraph.add_run(talk["venue"])
            set_run_font(venue, bold=True)
            paragraph.add_run(f", {talk['location']}, {talk['display_date']}.")
    participation = [
        talk for talk in talks if talk["presentation_type"] == "Conference participation"
    ]
    if participation:
        heading = doc.add_paragraph(style="CV Subsection")
        heading.add_run("Conference Participation")
        for talk in participation:
            paragraph = add_compact_paragraph(doc, left=0.23, hanging=0.23, after=1)
            venue = paragraph.add_run(talk["venue"])
            set_run_font(venue, bold=True)
            paragraph.add_run(f", {talk['location']}, {talk['display_date']}.")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor.from_string(TEXT)

    body = styles.add_style("CV Body", WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    body.font.name = BODY_FONT
    body.font.size = Pt(9.4)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(2)
    body.paragraph_format.line_spacing = 1.04
    body.paragraph_format.widow_control = True

    detail = styles.add_style("CV Detail", WD_STYLE_TYPE.PARAGRAPH)
    detail.base_style = body
    detail.font.color.rgb = RGBColor.from_string(MUTED)
    detail.paragraph_format.left_indent = Inches(0.18)
    detail.paragraph_format.space_after = Pt(1)

    section = styles.add_style("CV Section", WD_STYLE_TYPE.PARAGRAPH)
    section.base_style = body
    section.font.name = BODY_FONT
    section.font.size = Pt(10.8)
    section.font.bold = True
    section.font.color.rgb = RGBColor.from_string(NAVY)
    section.paragraph_format.space_before = Pt(8)
    section.paragraph_format.space_after = Pt(4)
    section.paragraph_format.keep_with_next = True

    subsection = styles.add_style("CV Subsection", WD_STYLE_TYPE.PARAGRAPH)
    subsection.base_style = body
    subsection.font.name = SERIF_FONT
    subsection.font.size = Pt(9.6)
    subsection.font.bold = True
    subsection.font.italic = True
    subsection.font.color.rgb = RGBColor.from_string(TEXT)
    subsection.paragraph_format.space_before = Pt(4)
    subsection.paragraph_format.space_after = Pt(1)
    subsection.paragraph_format.keep_with_next = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    header = section.header
    header.is_linked_to_previous = False
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    header_run = header_p.add_run("MENGFEI LI  ·  CURRICULUM VITAE")
    set_run_font(header_run, size=7.5, bold=True, color=MUTED)

    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    first_header.paragraphs[0].clear()

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    footer_run = footer_p.add_run("Mengfei Li  ·  ")
    set_run_font(footer_run, size=7.5, color=MUTED)
    add_page_field(footer_p)

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    first_footer_p = first_footer.paragraphs[0]
    first_footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_footer_p.paragraph_format.space_before = Pt(0)
    first_footer_run = first_footer_p.add_run("Mengfei Li  ·  ")
    set_run_font(first_footer_run, size=7.5, color=MUTED)
    add_page_field(first_footer_p)


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("MENGFEI LI")
    set_run_font(run, name=SERIF_FONT, size=22, bold=True, color=NAVY)

    identity = doc.add_paragraph()
    identity.alignment = WD_ALIGN_PARAGRAPH.CENTER
    identity.paragraph_format.space_after = Pt(2)
    run = identity.add_run("Ph.D. Candidate in Management Science  ·  Fudan University")
    set_run_font(run, size=10.2, bold=True, color=TEXT)

    market = doc.add_paragraph()
    market.alignment = WD_ALIGN_PARAGRAPH.CENTER
    market.paragraph_format.space_after = Pt(3)
    run = market.add_run(
        "2026-2027 ACADEMIC JOB MARKET  ·  OPERATIONS MANAGEMENT  ·  EXPECTED JUNE 2027"
    )
    set_run_font(run, size=8.2, bold=True, color=BLUE)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(7)
    add_hyperlink(contact, "mfli22@m.fudan.edu.cn", "mailto:mfli22@m.fudan.edu.cn")
    contact.add_run("  ·  ")
    add_hyperlink(contact, "Website", SITE_URL)
    contact.add_run("  ·  Shanghai, China")
    set_bottom_border(contact, color=LIGHT, size="6")


def build() -> Path:
    publications = [
        read_front_matter(path) for path in sorted((ROOT / "_publications").glob("*.md"))
    ]
    publications.sort(key=lambda item: item["sort_order"])
    talks = [read_front_matter(path) for path in sorted((ROOT / "_talks").glob("*.md"))]
    projects = parse_yaml((ROOT / "_data" / "projects.yml").read_text(encoding="utf-8"))
    awards = parse_yaml((ROOT / "_data" / "awards.yml").read_text(encoding="utf-8"))

    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    add_title_block(doc)

    add_section_heading(doc, "Education")
    add_date_entry(
        doc,
        "Ph.D. Candidate in Management Science, Fudan University",
        "Sep 2022 - Jun 2027 (expected)",
        [
            "School of Management, Department of Management Science",
            "Advisor: Prof. Xiaole Wu",
        ],
    )
    add_date_entry(
        doc,
        "Visiting Ph.D. Student, Indiana University",
        "Sep 2024 - Dec 2024",
        [
            "Kelley School of Business, Operations & Decision Technologies",
            "Host: Prof. Wenchang Zhang",
        ],
    )
    add_date_entry(
        doc,
        "B.S. in Statistics, Fudan University",
        "Sep 2018 - Jun 2022",
        [
            "School of Management, Department of Statistics",
            "Advisor: Prof. Feiyu Jiang",
        ],
    )

    add_section_heading(doc, "Research Interests")
    interests = [
        "AI-based Empirical Operations Management",
        "Causal Machine Learning",
        "Structural Model Estimation",
        "Manufacturing Learning, Efficiency, and Consistency",
        "Supply Chain Resilience",
    ]
    paragraph = add_compact_paragraph(doc, after=3)
    paragraph.add_run("  ·  ".join(interests))

    add_section_heading(doc, "Publications and Manuscripts")
    category_labels = [
        ("published-papers", "Published"),
        ("forthcoming", "Forthcoming"),
        ("under-review", "Under Review"),
        ("working-papers", "Working Papers"),
    ]
    for category, label in category_labels:
        items = [item for item in publications if item["category"] == category]
        if not items:
            continue
        subheading = doc.add_paragraph(style="CV Subsection")
        subheading.add_run(label)
        for number, publication in enumerate(items, start=1):
            add_publication(doc, publication, number)

    add_section_heading(doc, "Grants and Research Projects")
    for project in projects:
        add_project(doc, project)

    add_section_heading(doc, "Conference Presentations", page_break_before=True)
    add_talks(doc, talks, publications)

    add_section_heading(doc, "Honors and Awards")
    for award in awards:
        add_award(doc, award)

    add_section_heading(doc, "Teaching Experience")
    for item in (
        "Teaching Assistant, MBA course: Young Cadre Program in Science and Technology Innovation.",
        "Teaching Assistant, undergraduate course: Operations Management.",
    ):
        paragraph = add_compact_paragraph(doc, left=0.23, hanging=0.23, after=2)
        paragraph.add_run(item)

    add_section_heading(doc, "Academic Service")
    paragraph = add_compact_paragraph(doc, left=0.23, hanging=0.23, after=2)
    paragraph.add_run("Anonymous reviewer, ")
    journal = paragraph.add_run("Journal of Asian Economics")
    set_run_font(journal, name=SERIF_FONT, italic=True)
    paragraph.add_run(".")

    add_section_heading(doc, "Academic Resources and Software")
    activities = [
        "Participant, National Development and Intelligent Governance Laboratory Basic Research and Visualization Platform, Fudan University.",
        "Participant, China Foreign Merchandise Trade and Supply Chain Volatility Index, Fudan University Global Supply Chain Research Center.",
        "Contributor, Operations Management in the Digital Era (tentative title), edited by Xiaole Wu et al.",
    ]
    for item in activities:
        paragraph = add_compact_paragraph(doc, left=0.23, hanging=0.23, after=2)
        paragraph.add_run(item)
    paragraph = add_compact_paragraph(doc, left=0.23, hanging=0.23, after=2)
    paragraph.add_run("Creator and maintainer, R package ")
    add_hyperlink(paragraph, "SFHNV", "https://murphylicn.r-universe.dev/SFHNV")
    paragraph.add_run(".")

    properties = doc.core_properties
    properties.title = "Mengfei Li — Curriculum Vitae"
    properties.subject = "Academic curriculum vitae for Operations Management faculty recruitment"
    properties.author = "Mengfei Li"
    properties.last_modified_by = "Mengfei Li"
    properties.keywords = "Operations Management; Management Science; academic CV"
    properties.comments = ""
    properties.created = datetime(2026, 7, 27, tzinfo=timezone.utc)
    properties.modified = datetime(2026, 7, 27, tzinfo=timezone.utc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
